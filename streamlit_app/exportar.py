"""
Utilidades para exportar texto/markdown a PDF, DOCX y XLSX.
"""
import io
import re


def _limpiar_markdown(texto: str) -> str:
    """Elimina marcas Markdown básicas para texto plano."""
    texto = re.sub(r"#{1,6}\s+", "", texto)
    texto = re.sub(r"\*\*(.*?)\*\*", r"\1", texto)
    texto = re.sub(r"\*(.*?)\*", r"\1", texto)
    texto = re.sub(r"`(.*?)`", r"\1", texto)
    texto = re.sub(r"^[-*]\s+", "• ", texto, flags=re.MULTILINE)
    return texto.strip()


def a_pdf(texto: str, titulo: str = "Respuesta") -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    story = []

    # Título
    story.append(Paragraph(titulo, styles["Title"]))
    story.append(Spacer(1, 12))

    estilo_normal = ParagraphStyle(
        "custom", parent=styles["Normal"],
        fontSize=11, leading=16, spaceAfter=6,
    )
    estilo_h1 = ParagraphStyle(
        "h1", parent=styles["Heading1"], fontSize=14, spaceAfter=8,
    )
    estilo_h2 = ParagraphStyle(
        "h2", parent=styles["Heading2"], fontSize=12, spaceAfter=6,
    )

    for linea in texto.split("\n"):
        linea = linea.rstrip()
        if not linea:
            story.append(Spacer(1, 6))
        elif linea.startswith("## "):
            story.append(Paragraph(linea[3:], estilo_h1))
        elif linea.startswith("### "):
            story.append(Paragraph(linea[4:], estilo_h2))
        elif linea.startswith("# "):
            story.append(Paragraph(linea[2:], estilo_h1))
        elif linea.startswith("- ") or linea.startswith("* "):
            txt = "• " + linea[2:].replace("**", "").replace("*", "")
            story.append(Paragraph(txt, estilo_normal))
        else:
            # Convertir **negrita** a <b>
            linea_html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", linea)
            linea_html = re.sub(r"\*(.*?)\*", r"<i>\1</i>", linea_html)
            story.append(Paragraph(linea_html, estilo_normal))

    doc.build(story)
    return buf.getvalue()


def a_docx(texto: str, titulo: str = "Respuesta") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(titulo, level=0)

    for linea in texto.split("\n"):
        linea = linea.rstrip()
        if not linea:
            doc.add_paragraph("")
        elif linea.startswith("## ") or linea.startswith("# "):
            nivel = 1 if linea.startswith("# ") else 2
            doc.add_heading(linea.lstrip("#").strip(), level=nivel)
        elif linea.startswith("### "):
            doc.add_heading(linea[4:].strip(), level=3)
        elif linea.startswith("- ") or linea.startswith("* "):
            doc.add_paragraph(linea[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph()
            # Procesar negrita
            partes = re.split(r"\*\*(.*?)\*\*", linea)
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                if i % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def a_xlsx(texto: str, titulo: str = "Respuesta") -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    titulo_hoja = _re.sub(r'[\\/*?:\[\]]', '', titulo)[:30] or "Respuesta"
    ws.title = titulo_hoja

    ws.column_dimensions["A"].width = 120

    # Título
    ws["A1"] = titulo
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].fill = PatternFill("solid", fgColor="2E75B6")
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")

    fila = 3
    for linea in texto.split("\n"):
        linea = linea.rstrip()
        celda = ws.cell(row=fila, column=1, value=_limpiar_markdown(linea))
        celda.alignment = Alignment(wrap_text=True)

        if linea.startswith("## ") or linea.startswith("# "):
            celda.font = Font(bold=True, size=12)
            celda.fill = PatternFill("solid", fgColor="D5E8F0")
        elif linea.startswith("### "):
            celda.font = Font(bold=True, size=11)
        elif linea.startswith("- ") or linea.startswith("* "):
            celda.value = "  • " + _limpiar_markdown(linea[2:])

        fila += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()